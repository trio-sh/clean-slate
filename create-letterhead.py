"""
Generate professional Amani's Cleaners letterhead templates:
- Word (.docx) letterhead
- Excel (.xlsx) letterhead/invoice template
"""

import os
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
from openpyxl import Workbook
from openpyxl.styles import Font, Alignment, Border, Side, PatternFill, NamedStyle
from openpyxl.utils import get_column_letter
from openpyxl.drawing.image import Image as XlImage

LOGO_PATH = os.path.join(os.path.dirname(__file__), 'public', 'logo.png')
OUTPUT_DIR = os.path.join(os.path.dirname(__file__), 'public', 'marketing')

# Brand colors
NAVY = RGBColor(0x1a, 0x20, 0x55)
ORANGE = RGBColor(0xed, 0x6d, 0x1a)
GRAY = RGBColor(0x66, 0x66, 0x66)
LIGHT_GRAY = RGBColor(0x99, 0x99, 0x99)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)


# ─── WORD DOCUMENT ───────────────────────────────────────────────
def create_word_letterhead():
    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin = Cm(1.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    section = doc.sections[0]

    # ── HEADER ──
    header = section.header
    header.is_linked_to_previous = False

    # Header table: logo | company info | contact
    ht = header.add_table(rows=1, cols=3, width=Inches(6.5))
    ht.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Remove table borders
    for row in ht.rows:
        for cell in row.cells:
            cell._element.get_or_add_tcPr().append(
                parse_xml(f'<w:tcBorders {nsdecls("w")}>'
                          '<w:top w:val="none"/><w:left w:val="none"/>'
                          '<w:bottom w:val="none"/><w:right w:val="none"/>'
                          '</w:tcBorders>')
            )

    # Col 1: Logo
    cell_logo = ht.cell(0, 0)
    cell_logo.width = Inches(1.2)
    p_logo = cell_logo.paragraphs[0]
    p_logo.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run_logo = p_logo.add_run()
    run_logo.add_picture(LOGO_PATH, width=Inches(0.8))

    # Col 2: Company name
    cell_name = ht.cell(0, 1)
    cell_name.width = Inches(3.3)
    p_name = cell_name.paragraphs[0]
    p_name.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run_name = p_name.add_run("Amani's Cleaners")
    run_name.bold = True
    run_name.font.size = Pt(18)
    run_name.font.color.rgb = NAVY
    run_name.font.name = 'Georgia'

    p_sub = cell_name.add_paragraph()
    p_sub.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run_sub = p_sub.add_run("Premium Laundry & Dry Cleaning")
    run_sub.font.size = Pt(8)
    run_sub.font.color.rgb = ORANGE
    run_sub.font.name = 'Arial'
    run_sub.font.all_caps = True
    p_sub.paragraph_format.space_before = Pt(0)
    p_sub.paragraph_format.space_after = Pt(0)

    # Col 3: Contact info
    cell_contact = ht.cell(0, 2)
    cell_contact.width = Inches(2.0)

    contact_lines = [
        ("437-215-6321", False),
        ("amaniscleaners@gmail.com", False),
        ("amanicleaners.com", False),
        ("Toronto, ON, Canada", False),
    ]
    for i, (text, is_bold) in enumerate(contact_lines):
        if i == 0:
            p = cell_contact.paragraphs[0]
        else:
            p = cell_contact.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run(text)
        run.font.size = Pt(8)
        run.font.color.rgb = GRAY
        run.font.name = 'Arial'
        run.bold = is_bold

    # Orange line under header
    p_line = header.add_paragraph()
    p_line.paragraph_format.space_before = Pt(6)
    p_line.paragraph_format.space_after = Pt(0)
    # Add a bottom border via a horizontal rule
    p_line_pPr = p_line._element.get_or_add_pPr()
    pBdr = parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        '<w:bottom w:val="single" w:sz="8" w:space="1" w:color="ED6D1A"/>'
        '</w:pBdr>'
    )
    p_line_pPr.append(pBdr)

    # ── FOOTER ──
    footer = section.footer
    footer.is_linked_to_previous = False

    # Orange line above footer
    p_fline = footer.add_paragraph()
    p_fline.paragraph_format.space_before = Pt(0)
    p_fline.paragraph_format.space_after = Pt(4)
    pBdr_f = parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        '<w:top w:val="single" w:sz="8" w:space="1" w:color="ED6D1A"/>'
        '</w:pBdr>'
    )
    p_fline._element.get_or_add_pPr().append(pBdr_f)

    # Footer table: left info | center | right
    ft = footer.add_table(rows=1, cols=3, width=Inches(6.5))
    ft.alignment = WD_TABLE_ALIGNMENT.CENTER

    for row in ft.rows:
        for cell in row.cells:
            cell._element.get_or_add_tcPr().append(
                parse_xml(f'<w:tcBorders {nsdecls("w")}>'
                          '<w:top w:val="none"/><w:left w:val="none"/>'
                          '<w:bottom w:val="none"/><w:right w:val="none"/>'
                          '</w:tcBorders>')
            )

    # Footer left
    p_fl = ft.cell(0, 0).paragraphs[0]
    p_fl.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run_fl = p_fl.add_run("Amani's Cleaners  |  Proudly Canadian Since 2013")
    run_fl.font.size = Pt(7)
    run_fl.font.color.rgb = LIGHT_GRAY
    run_fl.font.name = 'Arial'

    # Footer center
    p_fc = ft.cell(0, 1).paragraphs[0]
    p_fc.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_fc = p_fc.add_run("amanicleaners.com")
    run_fc.font.size = Pt(7)
    run_fc.font.color.rgb = ORANGE
    run_fc.font.name = 'Arial'
    run_fc.bold = True

    # Footer right: page number
    p_fr = ft.cell(0, 2).paragraphs[0]
    p_fr.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run_fr = p_fr.add_run("Page ")
    run_fr.font.size = Pt(7)
    run_fr.font.color.rgb = LIGHT_GRAY
    run_fr.font.name = 'Arial'
    # Add page number field
    fldChar1 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="begin"/>')
    run_fr._element.append(fldChar1)
    instrText = parse_xml(f'<w:instrText {nsdecls("w")} xml:space="preserve"> PAGE </w:instrText>')
    run_fr._element.append(instrText)
    fldChar2 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="end"/>')
    run_fr._element.append(fldChar2)

    # ── BODY CONTENT (sample) ──
    # Add some space
    p_space = doc.add_paragraph()
    p_space.paragraph_format.space_before = Pt(12)

    # Date
    p_date = doc.add_paragraph()
    p_date.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run_date = p_date.add_run("[Date]")
    run_date.font.size = Pt(11)
    run_date.font.color.rgb = GRAY
    run_date.font.name = 'Arial'

    # Spacer
    doc.add_paragraph()

    # Recipient
    recipient_lines = [
        "[Recipient Name]",
        "[Company Name]",
        "[Street Address]",
        "[City, Province, Postal Code]",
    ]
    for line in recipient_lines:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run(line)
        run.font.size = Pt(11)
        run.font.color.rgb = GRAY
        run.font.name = 'Arial'

    # Spacer
    doc.add_paragraph()

    # Subject
    p_subj = doc.add_paragraph()
    run_subj = p_subj.add_run("RE: [Subject Line]")
    run_subj.bold = True
    run_subj.font.size = Pt(11)
    run_subj.font.color.rgb = NAVY
    run_subj.font.name = 'Arial'

    # Spacer
    doc.add_paragraph()

    # Dear
    p_dear = doc.add_paragraph()
    run_dear = p_dear.add_run("Dear [Recipient],")
    run_dear.font.size = Pt(11)
    run_dear.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    run_dear.font.name = 'Arial'

    doc.add_paragraph()

    # Body placeholder
    body_text = (
        "[Your letter content goes here. This is a professional letterhead template for "
        "Amani's Cleaners. Replace this text with your actual letter content. "
        "You can use this template for business correspondence, invoices, proposals, "
        "and any official communication.]"
    )
    p_body = doc.add_paragraph()
    p_body.paragraph_format.line_spacing = Pt(16)
    run_body = p_body.add_run(body_text)
    run_body.font.size = Pt(11)
    run_body.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    run_body.font.name = 'Arial'

    doc.add_paragraph()
    doc.add_paragraph()

    # Closing
    p_close = doc.add_paragraph()
    run_close = p_close.add_run("Sincerely,")
    run_close.font.size = Pt(11)
    run_close.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    run_close.font.name = 'Arial'

    doc.add_paragraph()
    doc.add_paragraph()

    # Signature block
    p_sig_name = doc.add_paragraph()
    run_sig = p_sig_name.add_run("[Your Name]")
    run_sig.bold = True
    run_sig.font.size = Pt(11)
    run_sig.font.color.rgb = NAVY
    run_sig.font.name = 'Arial'

    p_sig_title = doc.add_paragraph()
    p_sig_title.paragraph_format.space_before = Pt(0)
    run_title = p_sig_title.add_run("[Title]  |  Amani's Cleaners")
    run_title.font.size = Pt(10)
    run_title.font.color.rgb = GRAY
    run_title.font.name = 'Arial'

    p_sig_phone = doc.add_paragraph()
    p_sig_phone.paragraph_format.space_before = Pt(0)
    run_phone = p_sig_phone.add_run("437-215-6321  |  amaniscleaners@gmail.com")
    run_phone.font.size = Pt(9)
    run_phone.font.color.rgb = ORANGE
    run_phone.font.name = 'Arial'

    # Save
    path = os.path.join(OUTPUT_DIR, 'Amanis-Cleaners-Letterhead.docx')
    doc.save(path)
    print(f"Word letterhead saved: {path}")
    return path


# ─── EXCEL TEMPLATE ──────────────────────────────────────────────
def create_excel_letterhead():
    wb = Workbook()
    ws = wb.active
    ws.title = "Letterhead"

    # Colors
    navy_fill = PatternFill(start_color="1A2055", end_color="1A2055", fill_type="solid")
    orange_fill = PatternFill(start_color="ED6D1A", end_color="ED6D1A", fill_type="solid")
    light_gray_fill = PatternFill(start_color="F8F9FC", end_color="F8F9FC", fill_type="solid")
    white_fill = PatternFill(start_color="FFFFFF", end_color="FFFFFF", fill_type="solid")

    navy_font = Font(name='Arial', color="1A2055", bold=True)
    orange_font = Font(name='Arial', color="ED6D1A", bold=True)
    gray_font = Font(name='Arial', color="666666")
    white_font = Font(name='Arial', color="FFFFFF", bold=True)
    white_small = Font(name='Arial', color="FFFFFF", size=9)

    thin_border = Border(
        bottom=Side(style='thin', color='ED6D1A')
    )

    # Column widths
    ws.column_dimensions['A'].width = 3
    ws.column_dimensions['B'].width = 18
    ws.column_dimensions['C'].width = 18
    ws.column_dimensions['D'].width = 18
    ws.column_dimensions['E'].width = 18
    ws.column_dimensions['F'].width = 18
    ws.column_dimensions['G'].width = 3

    # Print setup
    ws.sheet_properties.pageSetUpPr = None
    ws.page_margins.left = 0.5
    ws.page_margins.right = 0.5
    ws.page_margins.top = 0.5
    ws.page_margins.bottom = 0.5

    # ── HEADER SECTION (Navy bar) ──
    for col in range(1, 8):
        cell = ws.cell(row=1, column=col)
        cell.fill = navy_fill

    for col in range(1, 8):
        cell = ws.cell(row=2, column=col)
        cell.fill = navy_fill

    for col in range(1, 8):
        cell = ws.cell(row=3, column=col)
        cell.fill = navy_fill

    ws.row_dimensions[1].height = 8
    ws.row_dimensions[2].height = 30
    ws.row_dimensions[3].height = 18

    # Add logo
    try:
        logo = XlImage(LOGO_PATH)
        logo.width = 45
        logo.height = 45
        ws.add_image(logo, 'B1')
    except Exception as e:
        print(f"Could not add logo to Excel: {e}")

    # Company name
    ws.merge_cells('C2:D2')
    cell_name = ws['C2']
    cell_name.value = "Amani's Cleaners"
    cell_name.font = Font(name='Georgia', size=16, bold=True, color="FFFFFF")
    cell_name.alignment = Alignment(vertical='center')
    cell_name.fill = navy_fill

    # Subtitle
    ws.merge_cells('C3:D3')
    cell_sub = ws['C3']
    cell_sub.value = "PREMIUM LAUNDRY & DRY CLEANING"
    cell_sub.font = Font(name='Arial', size=7, color="ED6D1A", bold=True)
    cell_sub.alignment = Alignment(vertical='top')
    cell_sub.fill = navy_fill

    # Contact info (right side)
    ws.merge_cells('E2:F2')
    cell_phone = ws['E2']
    cell_phone.value = "437-215-6321  |  amaniscleaners@gmail.com"
    cell_phone.font = white_small
    cell_phone.alignment = Alignment(horizontal='right', vertical='center')
    cell_phone.fill = navy_fill

    ws.merge_cells('E3:F3')
    cell_web = ws['E3']
    cell_web.value = "amanicleaners.com  |  Toronto, ON"
    cell_web.font = Font(name='Arial', size=8, color="F6B576")
    cell_web.alignment = Alignment(horizontal='right', vertical='top')
    cell_web.fill = navy_fill

    # Orange accent line
    ws.row_dimensions[4].height = 4
    for col in range(1, 8):
        cell = ws.cell(row=4, column=col)
        cell.fill = orange_fill

    # Spacer
    ws.row_dimensions[5].height = 20

    # ── DATE ──
    ws['B6'] = "[Date]"
    ws['B6'].font = gray_font
    ws['B6'].font = Font(name='Arial', size=11, color="666666")

    # Spacer
    ws.row_dimensions[7].height = 10

    # ── RECIPIENT ──
    recipient = [
        (8, "[Recipient Name]"),
        (9, "[Company Name]"),
        (10, "[Street Address]"),
        (11, "[City, Province, Postal Code]"),
    ]
    for row, text in recipient:
        ws.cell(row=row, column=2, value=text).font = Font(name='Arial', size=10, color="666666")
        ws.row_dimensions[row].height = 16

    ws.row_dimensions[12].height = 10

    # ── SUBJECT ──
    ws.merge_cells('B13:F13')
    ws['B13'] = "RE: [Subject Line]"
    ws['B13'].font = Font(name='Arial', size=11, bold=True, color="1A2055")

    ws.row_dimensions[14].height = 10

    # ── BODY ──
    ws.merge_cells('B15:F15')
    ws['B15'] = "Dear [Recipient],"
    ws['B15'].font = Font(name='Arial', size=10, color="333333")

    ws.row_dimensions[16].height = 10

    ws.merge_cells('B17:F19')
    body_cell = ws['B17']
    body_cell.value = (
        "[Your letter content goes here. This is a professional letterhead template for "
        "Amani's Cleaners. Replace this text with your actual letter content.]"
    )
    body_cell.font = Font(name='Arial', size=10, color="333333")
    body_cell.alignment = Alignment(wrap_text=True, vertical='top')

    ws.row_dimensions[17].height = 20
    ws.row_dimensions[18].height = 20
    ws.row_dimensions[19].height = 20

    ws.row_dimensions[20].height = 20
    ws.row_dimensions[21].height = 20

    # ── SIGNATURE ──
    ws['B22'] = "Sincerely,"
    ws['B22'].font = Font(name='Arial', size=10, color="333333")

    ws.row_dimensions[23].height = 25
    ws.row_dimensions[24].height = 25

    ws['B25'] = "[Your Name]"
    ws['B25'].font = Font(name='Arial', size=10, bold=True, color="1A2055")

    ws['B26'] = "[Title]  |  Amani's Cleaners"
    ws['B26'].font = Font(name='Arial', size=9, color="666666")

    ws['B27'] = "437-215-6321  |  amaniscleaners@gmail.com"
    ws['B27'].font = Font(name='Arial', size=9, color="ED6D1A")

    # ── FOOTER ──
    # Find a row near bottom
    footer_row = 38

    ws.row_dimensions[footer_row - 1].height = 4
    for col in range(1, 8):
        ws.cell(row=footer_row - 1, column=col).fill = orange_fill

    for col in range(1, 8):
        ws.cell(row=footer_row, column=col).fill = navy_fill

    ws.row_dimensions[footer_row].height = 22

    ws.merge_cells(f'B{footer_row}:C{footer_row}')
    footer_left = ws.cell(row=footer_row, column=2)
    footer_left.value = "Amani's Cleaners  |  Proudly Canadian Since 2013"
    footer_left.font = Font(name='Arial', size=7, color="999999")
    footer_left.alignment = Alignment(vertical='center')
    footer_left.fill = navy_fill

    ws.merge_cells(f'D{footer_row}:D{footer_row}')
    footer_center = ws.cell(row=footer_row, column=4)
    footer_center.value = "amanicleaners.com"
    footer_center.font = Font(name='Arial', size=7, color="ED6D1A", bold=True)
    footer_center.alignment = Alignment(horizontal='center', vertical='center')
    footer_center.fill = navy_fill

    ws.merge_cells(f'E{footer_row}:F{footer_row}')
    footer_right = ws.cell(row=footer_row, column=5)
    footer_right.value = "437-215-6321  |  amaniscleaners@gmail.com"
    footer_right.font = Font(name='Arial', size=7, color="999999")
    footer_right.alignment = Alignment(horizontal='right', vertical='center')
    footer_right.fill = navy_fill

    # Print area
    ws.print_area = 'A1:G38'

    # ── INVOICE SHEET ──
    ws2 = wb.create_sheet("Invoice Template")

    ws2.column_dimensions['A'].width = 3
    ws2.column_dimensions['B'].width = 28
    ws2.column_dimensions['C'].width = 12
    ws2.column_dimensions['D'].width = 12
    ws2.column_dimensions['E'].width = 15
    ws2.column_dimensions['F'].width = 15
    ws2.column_dimensions['G'].width = 3

    # Header bar
    for r in range(1, 4):
        for col in range(1, 8):
            ws2.cell(row=r, column=col).fill = navy_fill

    ws2.row_dimensions[1].height = 8
    ws2.row_dimensions[2].height = 30
    ws2.row_dimensions[3].height = 18

    try:
        logo2 = XlImage(LOGO_PATH)
        logo2.width = 45
        logo2.height = 45
        ws2.add_image(logo2, 'B1')
    except:
        pass

    ws2.merge_cells('C2:D2')
    ws2['C2'].value = "Amani's Cleaners"
    ws2['C2'].font = Font(name='Georgia', size=16, bold=True, color="FFFFFF")
    ws2['C2'].alignment = Alignment(vertical='center')
    ws2['C2'].fill = navy_fill

    ws2.merge_cells('C3:D3')
    ws2['C3'].value = "PREMIUM LAUNDRY & DRY CLEANING"
    ws2['C3'].font = Font(name='Arial', size=7, color="ED6D1A", bold=True)
    ws2['C3'].fill = navy_fill

    ws2.merge_cells('E2:F2')
    ws2['E2'].value = "INVOICE"
    ws2['E2'].font = Font(name='Georgia', size=18, bold=True, color="ED6D1A")
    ws2['E2'].alignment = Alignment(horizontal='right', vertical='center')
    ws2['E2'].fill = navy_fill

    # Orange line
    ws2.row_dimensions[4].height = 4
    for col in range(1, 8):
        ws2.cell(row=4, column=col).fill = orange_fill

    ws2.row_dimensions[5].height = 14

    # Invoice details
    ws2['B6'] = "Bill To:"
    ws2['B6'].font = Font(name='Arial', size=8, color="999999", bold=True)
    ws2['E6'] = "Invoice #:"
    ws2['E6'].font = Font(name='Arial', size=9, color="666666")
    ws2['F6'] = "[INV-0001]"
    ws2['F6'].font = Font(name='Arial', size=9, color="1A2055", bold=True)
    ws2['F6'].alignment = Alignment(horizontal='right')

    ws2['B7'] = "[Client Name]"
    ws2['B7'].font = Font(name='Arial', size=10, color="333333", bold=True)
    ws2['E7'] = "Date:"
    ws2['E7'].font = Font(name='Arial', size=9, color="666666")
    ws2['F7'] = "[DD/MM/YYYY]"
    ws2['F7'].font = Font(name='Arial', size=9, color="1A2055")
    ws2['F7'].alignment = Alignment(horizontal='right')

    ws2['B8'] = "[Client Address]"
    ws2['B8'].font = Font(name='Arial', size=9, color="666666")
    ws2['E8'] = "Due Date:"
    ws2['E8'].font = Font(name='Arial', size=9, color="666666")
    ws2['F8'] = "[DD/MM/YYYY]"
    ws2['F8'].font = Font(name='Arial', size=9, color="ED6D1A", bold=True)
    ws2['F8'].alignment = Alignment(horizontal='right')

    ws2['B9'] = "[City, Province, Postal Code]"
    ws2['B9'].font = Font(name='Arial', size=9, color="666666")

    ws2.row_dimensions[10].height = 10

    # Table header
    headers = ['Service / Item', 'Qty', 'Unit Price', 'Amount']
    cols = [2, 4, 5, 6]
    header_fill = PatternFill(start_color="1A2055", end_color="1A2055", fill_type="solid")

    ws2.merge_cells('B11:C11')
    ws2['B11'] = headers[0]
    ws2['B11'].font = Font(name='Arial', size=9, bold=True, color="FFFFFF")
    ws2['B11'].fill = header_fill
    ws2['B11'].alignment = Alignment(vertical='center')
    ws2['C11'].fill = header_fill

    for i, (h, c) in enumerate(zip(headers[1:], [4, 5, 6])):
        cell = ws2.cell(row=11, column=c, value=h)
        cell.font = Font(name='Arial', size=9, bold=True, color="FFFFFF")
        cell.fill = header_fill
        cell.alignment = Alignment(horizontal='right' if c > 3 else 'center', vertical='center')

    ws2.row_dimensions[11].height = 28

    # Sample rows
    alt_fill = PatternFill(start_color="F8F9FC", end_color="F8F9FC", fill_type="solid")
    for row in range(12, 20):
        ws2.merge_cells(f'B{row}:C{row}')
        ws2.row_dimensions[row].height = 22

        item_font = Font(name='Arial', size=9, color="333333")
        num_font = Font(name='Arial', size=9, color="333333")

        if row == 12:
            ws2[f'B{row}'] = "[Service name]"
            ws2[f'D{row}'] = "[qty]"
            ws2[f'E{row}'] = "[$0.00]"
            ws2[f'F{row}'] = "[$0.00]"
        else:
            ws2[f'B{row}'] = ""

        ws2[f'B{row}'].font = item_font
        ws2[f'D{row}'].font = num_font
        ws2[f'D{row}'].alignment = Alignment(horizontal='center')
        ws2[f'E{row}'].font = num_font
        ws2[f'E{row}'].alignment = Alignment(horizontal='right')
        ws2[f'F{row}'].font = num_font
        ws2[f'F{row}'].alignment = Alignment(horizontal='right')

        if row % 2 == 0:
            for col in range(2, 7):
                ws2.cell(row=row, column=col).fill = alt_fill

    # Totals
    ws2.row_dimensions[20].height = 6

    totals = [
        (21, "Subtotal:", "[$0.00]", False),
        (22, "HST (13%):", "[$0.00]", False),
        (23, "Total:", "[$0.00]", True),
    ]

    for row, label, value, is_total in totals:
        ws2[f'E{row}'] = label
        ws2[f'E{row}'].font = Font(name='Arial', size=9, color="666666", bold=is_total)
        ws2[f'E{row}'].alignment = Alignment(horizontal='right')

        ws2[f'F{row}'] = value
        if is_total:
            ws2[f'F{row}'].font = Font(name='Arial', size=12, color="1A2055", bold=True)
            ws2[f'F{row}'].fill = PatternFill(start_color="FEF7F0", end_color="FEF7F0", fill_type="solid")
        else:
            ws2[f'F{row}'].font = Font(name='Arial', size=9, color="333333")
        ws2[f'F{row}'].alignment = Alignment(horizontal='right')

    # Notes
    ws2.row_dimensions[25].height = 10
    ws2['B26'] = "Notes / Terms:"
    ws2['B26'].font = Font(name='Arial', size=8, color="999999", bold=True)
    ws2.merge_cells('B27:F28')
    ws2['B27'] = "[Payment terms, special instructions, or notes]"
    ws2['B27'].font = Font(name='Arial', size=9, color="666666")
    ws2['B27'].alignment = Alignment(wrap_text=True, vertical='top')

    # Footer
    footer_row2 = 32
    ws2.row_dimensions[footer_row2 - 1].height = 4
    for col in range(1, 8):
        ws2.cell(row=footer_row2 - 1, column=col).fill = orange_fill

    for col in range(1, 8):
        ws2.cell(row=footer_row2, column=col).fill = navy_fill
    ws2.row_dimensions[footer_row2].height = 22

    ws2.merge_cells(f'B{footer_row2}:C{footer_row2}')
    ws2[f'B{footer_row2}'].value = "Thank you for choosing Amani's Cleaners!"
    ws2[f'B{footer_row2}'].font = Font(name='Arial', size=7, color="F6B576")
    ws2[f'B{footer_row2}'].alignment = Alignment(vertical='center')
    ws2[f'B{footer_row2}'].fill = navy_fill

    ws2.merge_cells(f'E{footer_row2}:F{footer_row2}')
    ws2[f'E{footer_row2}'].value = "437-215-6321  |  amanicleaners.com"
    ws2[f'E{footer_row2}'].font = Font(name='Arial', size=7, color="999999")
    ws2[f'E{footer_row2}'].alignment = Alignment(horizontal='right', vertical='center')
    ws2[f'E{footer_row2}'].fill = navy_fill

    # Save
    path = os.path.join(OUTPUT_DIR, 'Amanis-Cleaners-Letterhead.xlsx')
    wb.save(path)
    print(f"Excel letterhead saved: {path}")
    return path


if __name__ == '__main__':
    os.makedirs(OUTPUT_DIR, exist_ok=True)
    create_word_letterhead()
    create_excel_letterhead()
    print("\nDone! Files are in public/marketing/")
